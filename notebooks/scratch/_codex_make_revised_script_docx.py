from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path.cwd()
OUT = ROOT / "docs" / "final_presentation_script_revised.docx"
FALLBACK_OUT = ROOT / "docs" / "final_presentation_script_revised_v2.docx"


SCRIPT = [
    (
        "1. 제목",
        "안녕하세요. 저희는 BOSCH 플라즈마 식각 공정에서 OES와 process sensor 데이터를 이용해, 식각 후 웨이퍼 계측값을 예측하는 Virtual Metrology 모델을 만들었습니다. 데이터와 baseline은 짧게 설명드리고, 뒤쪽의 cycle-aware 아키텍처와 개선 과정에 집중해서 말씀드리겠습니다.",
    ),
    (
        "2. Motivation",
        "반도체 공정에서 웨이퍼 품질은 식각 깊이나 산화막 두께 같은 물리적 계측값으로 판단합니다. 하지만 직접 계측은 시간과 비용이 크기 때문에 모든 웨이퍼에 적용하기 어렵습니다. VM은 공정 중 센서 데이터만으로 이 계측값을 예측해 검사 비용을 줄이고, 공정 제어를 빠르게 만드는 기술입니다.",
    ),
    (
        "3. VM 선행연구",
        "기존 VM 연구도 OES 데이터를 많이 사용했지만, 평균이나 분산, PCA band처럼 통계 feature로 압축하는 경우가 많았습니다. 딥러닝 연구가 나오면서 시간-파장 구조를 더 잘 활용하게 되었지만, BOSCH 공정의 etch와 passivation 반복 cycle을 직접 반영하는 데는 한계가 있었습니다. 저희는 이 cycle 구조를 모델 안에 넣는 방향으로 접근했습니다.",
    ),
    (
        "4. Dataset",
        "데이터는 실제 BOSCH plasma etching 장비에서 수집된 공개 데이터입니다. 200mm silicon wafer에 SF6 etch와 C4F8 passivation을 반복하는 100 cycle 공정이 적용되어 있습니다. 입력은 3,648개 OES wavelength와 31개 공정 sensor이고, 정답값은 식각 후 측정한 si_etch와 oxide_etch입니다.",
    ),
    (
        "5. Target 및 데이터 파이프라인",
        "모델이 맞춰야 하는 값은 식각이 끝난 뒤 실제로 측정되는 etch 값입니다. 학습할 때는 공정 중 기록된 OES, process sensor, 측정 위치 XY를 입력으로 넣고, 식각 후 계측값을 label로 사용합니다. 목표는 나중에 실제 계측 없이도 센서 데이터만으로 식각 결과를 예측하는 것입니다.",
    ),
    (
        "6. XGBoost Baseline",
        "먼저 비교 기준으로 XGBoost baseline을 만들었습니다. OES와 process signal을 cycle 단위로 요약하고 평균, 표준편차, 최솟값, 최댓값 같은 통계 feature를 뽑은 뒤 XY 좌표를 붙였습니다. 이 baseline은 딥러닝 모델이 정말 의미 있는 개선을 만드는지 판단하기 위한 기준선입니다.",
    ),
    (
        "7. 초기 Cycle-Aware DL",
        "다음은 초기 DL 모델입니다. BOSCH 공정은 100개의 cycle이 반복되기 때문에, 전체 신호를 평균내기보다 cycle별 패턴과 cycle 간 변화를 따로 보는 것이 중요하다고 봤습니다. 그래서 cycle 내부는 CNN encoder가, cycle 사이의 흐름은 BiLSTM이 학습하도록 설계했습니다.",
    ),
    (
        "8. 딥러닝 구조 정리",
        "초기에는 XGBoost와 딥러닝 ensemble까지 고려했지만, 최종적으로는 단일 딥러닝 구조가 baseline보다 나아지는지 먼저 검증하는 방향으로 정리했습니다. 이때 슬라이드 7의 구조도 단순화해서, OES와 공정 파라미터를 서로 다른 인코더로 두지 않고 같은 2D-CNN 아키텍처로 처리하되, weight는 모달리티별로 독립적으로 학습하도록 했습니다. 이후 실험은 OES와 process sensor를 함께 쓰는 multimodal 딥러닝을 중심으로 진행했습니다.",
    ),
    (
        "9. 문제 1: DL이 Si를 못 맞추는 문제",
        "그런데 처음에는 예상과 다른 문제가 있었습니다. XGBoost는 si_etch를 매우 잘 맞추는데, 초기 DL은 오히려 Si 성능이 크게 떨어졌습니다. 모델이 더 복잡하다고 항상 좋아지는 것이 아니라, 중요한 정보를 모델이 어떻게 받아들이는지가 더 중요하다는 뜻이었습니다.",
    ),
    (
        "10. EDA - Si_etch",
        "원인을 보려고 EDA를 해보니, si_etch는 웨이퍼 안에서 위치에 따라 달라지는 spatial pattern이 매우 강했습니다. 표에서 보시듯 Si는 wafer 간 차이보다 wafer 내부 위치 차이가 훨씬 크고, XY만 써도 높은 설명력이 나옵니다. XGBoost는 XY feature를 직접 활용하기 쉬웠지만, 초기 DL은 wafer representation을 만든 뒤 마지막에 XY를 붙이는 정도라 위치 정보가 충분히 살아나지 못했습니다.",
    ),
    (
        "11. FiLM + Fourier XY를 적용한 DL",
        "그래서 구조를 바꿨습니다. OES와 process branch로 wafer-level representation을 만든 뒤, 측정 위치인 XY 좌표를 Fourier feature로 변환했습니다. 그리고 이 XY 정보를 단순히 마지막에 concat하는 대신, FiLM modulation으로 wafer representation 자체를 위치별로 조절하게 만들었습니다. 같은 wafer라도 중심부와 edge는 다른 point-wise representation을 갖게 되는 방식입니다.",
    ),
    (
        "12. FiLM modulation",
        "기존 concat 방식은 wafer feature가 이미 만들어진 뒤 XY가 뒤늦게 붙는 구조입니다. 이 경우 OES와 process에서 온 큰 feature에 비해 XY의 영향이 약해질 수 있습니다. 반면 FiLM은 XY로부터 gamma와 beta를 만들어 feature를 scaling하고 shifting합니다. 쉽게 말해 위치 정보가 모델 안쪽에서 feature를 직접 조절하도록 만든 것입니다.",
    ),
    (
        "13. FiLM 적용 결과",
        "이 구조를 적용한 뒤에는 초기 DL에서 무너졌던 Si 예측 성능이 크게 회복되었습니다. 이 결과는 Si 문제의 원인이 데이터 부족이라기보다는, spatial conditioning을 충분히 강하게 넣지 못했던 모델 구조의 문제였다는 점을 보여줍니다. 그래서 이후 모델에서는 Fourier XY와 FiLM을 기본 구조로 유지했습니다.",
    ),
    (
        "14. Si_etch wafer-map",
        "wafer map으로 보면 차이가 더 직관적입니다. 초기 DL은 위치별 Si pattern을 제대로 따라가지 못했지만, FiLM과 Fourier XY를 적용한 모델은 실제 wafer map의 분포를 훨씬 잘 재현했습니다. Si는 이 단계에서 충분히 회복되었기 때문에, 이후에는 더 어려운 oxide_etch 성능을 안정적으로 끌어올리는 데 집중했습니다.",
    ),
    (
        "15. 문제 2: 특정 Fold에서 성능 붕괴",
        "oxide_etch로 넘어오면서 두 번째 문제가 나타났습니다. 단일 fold에서는 성능이 좋아 보였지만, 5-fold로 확장하자 특정 fold에서 성능이 크게 무너졌습니다. 특히 초기 DL은 fold 0과 3에서는 R2가 0.74, 0.75 수준까지 나오지만, fold 2는 0.415, fold 4는 0.149까지 떨어졌습니다. 평균 성능만 보면 놓칠 수 있는 불안정성이 여기서 드러났습니다.",
    ),
    (
        "16. 원인 분석",
        "원인을 보면 fold 2와 fold 4의 문제가 조금 달랐습니다. fold 4는 실제 wafer mean은 넓게 퍼져 있는데 예측값은 0.64에서 0.66 근처로 눌리는 collapse가 보였습니다. predicted spread가 true spread의 0.19배밖에 되지 않았습니다. 반면 fold 2는 분산 자체가 완전히 죽은 것은 아니지만, validation RMSE가 초반에 가장 좋고 이후 흔들리는 early overfitting 경향이 강했습니다. 그래서 feature noise를 줄이는 방법과 학습을 안정화하는 방법을 같이 넣어야겠다고 판단했습니다.",
    ),
    (
        "17. Improve 1: OES top-k",
        "첫 번째 개선은 OES wavelength selection입니다. OES는 3,648개 wavelength를 모두 쓰기 때문에, 작은 데이터셋에서는 target과 거의 관련 없는 wavelength가 noise로 작동할 수 있습니다. 그래서 각 fold의 train set에서만 oxide_etch와 wavelength summary statistic의 correlation을 계산하고, 상관성이 높은 상위 256개 wavelength만 남겼습니다. 여기서 중요한 점은 validation 정보를 보지 않고 train fold 안에서만 선택했다는 것입니다.",
    ),
    (
        "18. Improve 2: Auxiliary wafer-mean loss",
        "두 번째는 auxiliary wafer-mean loss입니다. 기존 모델은 89개 point의 값을 각각 맞추는 데 집중했기 때문에, wafer 전체가 high etch인지 low etch인지 같은 global level을 representation에 충분히 담지 못할 수 있었습니다. 그래서 wafer representation만 보고 wafer 평균 oxide 값을 예측하는 보조 head를 추가했습니다. point prediction loss에 wafer mean loss를 같이 더해서, 모델이 point별 차이뿐 아니라 wafer 간 level 차이도 놓치지 않도록 한 것입니다.",
    ),
    (
        "19. Improve 3: Mixup",
        "세 번째는 wafer-level mixup입니다. 레이블이 있는 wafer 수가 많지 않기 때문에, 모델이 특정 wafer 조합에 과하게 맞춰질 위험이 있었습니다. 그래서 두 wafer의 OES tensor, process tensor, XY feature, target을 같은 비율로 섞어 중간 상태의 wafer도 학습하도록 했습니다. 이렇게 하면 모델이 train set의 몇 개 wafer만 외우는 대신, wafer 사이의 연속적인 변화를 더 부드럽게 보게 됩니다.",
    ),
    (
        "20. Improve 4: Weight EMA",
        "네 번째는 weight EMA입니다. 학습 중 validation RMSE가 흔들리면, 우연히 한 번 낮게 나온 epoch가 best checkpoint로 저장될 수 있습니다. EMA는 현재 weight만 쓰지 않고, 학습 과정에서 지나온 weight를 지수이동평균으로 누적한 shadow model을 사용합니다. 최근 weight를 더 크게 반영하면서도 전체 trajectory를 부드럽게 만들어, fold별 성능 변동을 줄이는 효과를 기대했습니다.",
    ),
    (
        "21. Final Pipeline",
        "최종 파이프라인은 이렇게 정리됩니다. 아키텍처 쪽에서는 FiLM + Fourier XY를 유지하고, oxide 예측을 위해 OES top-k selection과 auxiliary wafer-mean head를 추가했습니다. 학습 과정에서는 wafer-level mixup과 EMA를 적용했습니다. 정리하면, 앞의 두 가지는 모델이 어떤 정보를 보게 할지에 대한 개선이고, 뒤의 두 가지는 작은 데이터에서 학습이 흔들리지 않게 하는 개선입니다.",
    ),
    (
        "22. Results - Ablation",
        "개선 효과는 ablation으로 확인했습니다. XGBoost baseline의 oxide R2는 평균 0.551이고, 초기 DL 5-fold는 0.543으로 거의 비슷했습니다. 여기에 OES top-k를 넣으면 0.596, auxiliary mean loss를 넣으면 0.621, mixup까지 적용하면 0.644로 올라갔습니다. 마지막으로 EMA를 적용한 최종 모델은 평균 R2 0.666까지 개선되었습니다. 한 번에 좋아진 것이 아니라, 각 개선이 조금씩 안정성을 올린 흐름입니다.",
    ),
    (
        "23. Results - Fold별 변화",
        "fold별로 보면 개선 이유가 더 분명합니다. 초기 DL에서 가장 크게 무너졌던 fold 4는 R2 0.15에서 최종 0.59까지 회복되었습니다. fold 2도 0.41에서 0.49로 올라갔고, fold 0, 1, 3은 0.72에서 0.78 수준으로 안정적으로 유지되었습니다. 아직 모든 fold가 완전히 균일하다고 보기는 어렵지만, 특정 fold 붕괴는 상당히 완화되었습니다.",
    ),
    (
        "24. Results - Error analysis",
        "잔차 분포를 보면 final DL의 residual이 0 근처에 모여 있고, XGBoost보다 조금 더 좁게 모입니다. 다만 fold 2와 fold 4에서는 error spread가 여전히 큽니다. 전체적으로는 개선됐지만, 어려운 validation fold가 남아 있다는 점은 한계로 봐야 합니다.",
    ),
    (
        "25. Results - Spatial error map",
        "spatial error map은 예측 오차가 특정 위치에 몰리는지 확인하기 위한 시각화입니다. 평균 절대오차가 큰 point가 일부 있지만, 한 방향으로 계속 과대 또는 과소 예측하는 강한 위치 편향은 크지 않았습니다. 남은 문제는 단순한 XY 문제가 아니라, wafer 또는 lot 단위 상태 차이를 더 잘 일반화하는 문제에 가깝다고 봤습니다.",
    ),
    (
        "26. Results - 5-fold Robustness",
        "5-fold 관점에서 보면 단일 fold 결과만 보고 판단하면 위험하다는 점이 드러납니다. single-fold에서는 R2 0.734까지 나왔지만, 초기 5-fold 평균은 0.543으로 크게 떨어졌습니다. 최종 모델은 평균 0.666까지 올라가서 single-fold보다는 낮지만, 여러 fold에서 더 안정적인 성능을 보였습니다. 그래서 최종 성능은 5-fold 기준으로 보는 것이 더 보수적이고 타당하다고 판단했습니다.",
    ),
    (
        "27. Results - XGBoost vs Final DL",
        "XGBoost와 최종 DL을 fold별로 비교하면, final DL은 fold 0, 1, 2, 3에서 XGBoost보다 높은 R2를 보였습니다. fold 4에서는 XGBoost가 0.62, final DL이 0.59로 조금 낮습니다. 그래서 이 결과는 DL이 모든 fold에서 무조건 이겼다는 의미보다는, 대부분의 fold에서 baseline을 넘었고 동시에 fold 4 같은 hard case는 아직 개선 여지가 남아 있다는 의미로 해석했습니다.",
    ),
    (
        "28. Results - 전체 validation point 비교",
        "마지막으로 전체 validation point를 모아 보면 차이가 더 분명합니다. XGBoost baseline은 pooled R2가 0.556, RMSE가 0.0515였고, final DL은 pooled R2 0.668, RMSE 0.0445까지 개선되었습니다. 산점도에서도 final DL 쪽이 대각선에 더 가깝게 모여 있습니다. 결론적으로 cycle-aware multimodal 구조와 네 가지 개선을 함께 적용했을 때, oxide_etch 예측에서도 baseline 대비 의미 있는 향상을 확인했습니다.",
    ),
    (
        "29. Limitation - Lot 단위 강건성",
        "다만 현장 적용 관점에서는 한계가 있습니다. 일반 5-fold에서는 최종 모델이 R2 0.666까지 나오지만, 특정 lot을 통째로 제외하고 처음 보는 lot처럼 예측하면 R2가 0.323까지 떨어졌습니다. 날짜나 lot이 달라지면 장비 상태와 chamber condition, sensor drift가 함께 바뀔 수 있기 때문에, lot-level generalization은 아직 충분하지 않습니다.",
    ),
    (
        "30. Limitation 및 개선 방향",
        "개선 방향은 네 가지입니다. 첫째, 더 다양한 lot과 공정 조건의 labeled data가 필요합니다. 둘째, BiLSTM뿐 아니라 Transformer나 TCN처럼 cycle 중요도를 더 직접적으로 볼 수 있는 encoder를 비교해야 합니다. 셋째, 실제 운영을 위해 inference time과 이상 wafer 탐지 연계를 검증해야 합니다. 마지막으로, 어떤 wavelength, cycle, sensor가 예측에 영향을 줬는지 최종 모델 전체 fold 기준으로 해석 분석을 확장해야 합니다.",
    ),
    (
        "31. Expected Impact - 공정 해석",
        "기대 효과는 값을 하나 예측하는 데서 끝나지 않습니다. 모델이 입력을 cycle 단위로 보기 때문에, 나중에 attribution이나 attention 해석을 붙이면 어떤 cycle 구간에서 이상 신호가 커졌는지 추적할 수 있습니다. VM을 품질 예측 도구뿐 아니라 공정 drift 해석 도구로도 활용할 수 있습니다.",
    ),
    (
        "32. Expected Impact - 확장성",
        "마지막으로 이 접근은 BOSCH 식각 공정에만 한정되지 않습니다. ALD, PEALD, ALE처럼 step이 반복되는 cycle 공정에서도 cycle별 sensor pattern과 cycle 간 drift가 중요합니다. 따라서 이번 cycle-aware 구조는 다른 반복 공정의 VM으로도 확장할 수 있습니다. 이상으로 발표 마치겠습니다. 감사합니다.",
    ),
]


def set_run_font(run, name: str = "Malgun Gothic", size_pt: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("최종발표 대본 수정본")
    title_run.bold = True
    set_run_font(title_run, size_pt=16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("10분 내외 발표 기준: 초반 압축, 아키텍처와 개선 과정 중심")
    set_run_font(subtitle_run, size_pt=10)
    subtitle.paragraph_format.space_after = Pt(10)

    for idx, (heading, body) in enumerate(SCRIPT, 1):
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(5 if idx > 1 else 0)
        hp.paragraph_format.space_after = Pt(1)
        hr = hp.add_run(heading)
        hr.bold = True
        set_run_font(hr, size_pt=11)

        bp = doc.add_paragraph()
        bp.paragraph_format.line_spacing = 1.12
        bp.paragraph_format.space_after = Pt(4)
        br = bp.add_run(body)
        set_run_font(br, size_pt=10.5)

    try:
        doc.save(OUT)
        saved = OUT
    except PermissionError:
        doc.save(FALLBACK_OUT)
        saved = FALLBACK_OUT
    print(saved)


if __name__ == "__main__":
    main()
